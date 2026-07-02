import openpyxl
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_excel(file_path):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheet = wb.active

    header_row = [cell.value for cell in sheet[1]]
    age_row = [cell.value for cell in sheet[2]]

    grade_names = ['پیش‌دبستانی', 'دبستان1', 'دبستان2', 'متوسطه1', 'متوسطه2', 'بزرگسالان']
    col_indices = list(range(1, 7))

    institutes = []
    current_institute = None
    row_idx = 2
    total_rows = sheet.max_row

    def is_institute_row(row):
        first = row[0].value
        if not first:
            return False
        first = str(first).strip()
        keywords = ['درصد', 'شماره تماس', 'تلفن', 'آدرس']
        if any(k in first for k in keywords):
            return True
        if (not re.search(r'\d', first)) and len(first) > 3 and re.search(r'[آ-ی]', first):
            return True
        return False

    def is_course_row(row):
        for col in col_indices:
            val = row[col].value
            if val and str(val).strip():
                if not re.search(r'تومان', str(val)) and not isinstance(val, (int, float)):
                    return True
        return False

    def is_price_row(row):
        for col in col_indices:
            val = row[col].value
            if val:
                val_str = str(val).strip()
                if 'تومان' in val_str or re.search(r'\d', val_str):
                    return True
        return False

    def parse_institute(row):
        first = str(row[0].value).strip() if row[0].value else ''
        name = first
        discount = ''
        address = ''
        phone = ''

        disc_match = re.search(r'(\d+)\s*درصد', first)
        if disc_match:
            discount = disc_match.group(1) + '%'
            name = first.replace(disc_match.group(0), '').strip()

        if 'آدرس' in first:
            address = first

        phone_match = re.search(r'۰?۹[۰-۹]{9}', first)
        if phone_match:
            phone = phone_match.group(0)

        return {'name': name, 'discount': discount, 'address': address, 'phone': phone}

    while row_idx < total_rows:
        row = sheet[row_idx]
        if all(cell.value is None or str(cell.value).strip() == '' for cell in row):
            row_idx += 1
            continue

        if is_institute_row(row):
            if current_institute:
                institutes.append(current_institute)

            parsed = parse_institute(row)
            current_institute = {
                'name': parsed['name'],
                'discount': parsed['discount'],
                'address': parsed['address'],
                'phone': parsed['phone'],
                'courses': []
            }

            row_idx += 1
            course_rows = []
            price_rows = []

            while row_idx < total_rows:
                next_row = sheet[row_idx]
                if is_institute_row(next_row):
                    break

                if is_course_row(next_row):
                    course_rows.append(next_row)
                elif is_price_row(next_row):
                    if course_rows:
                        price_rows.append(next_row)
                row_idx += 1

            while len(price_rows) < len(course_rows):
                price_rows.append(None)

            for course_row, price_row in zip(course_rows, price_rows):
                for col_idx in col_indices:
                    course_name = course_row[col_idx].value
                    if not course_name:
                        continue
                    course_name = str(course_name).strip()

                    price_val = None
                    price_text = ''
                    if price_row:
                        price_cell = price_row[col_idx].value
                        if price_cell:
                            price_text = str(price_cell).strip()
                            num_match = re.search(r'([\d,]+)', price_text)
                            if num_match:
                                num_str = num_match.group(1).replace(',', '')
                                for fa, en in zip('۰۱۲۳۴۵۶۷۸۹', '0123456789'):
                                    num_str = num_str.replace(fa, en)
                                try:
                                    price_val = float(num_str)
                                except ValueError:
                                    price_val = None

                    grade_index = col_idx - 1
                    if grade_index < len(grade_names):
                        grade = grade_names[grade_index]
                    else:
                        grade = f'مقطع{col_idx}'

                    current_institute['courses'].append({
                        'grade': grade,
                        'course': course_name,
                        'price': price_val,
                        'price_text': price_text
                    })

            continue
        else:
            row_idx += 1

    if current_institute:
        institutes.append(current_institute)

    return institutes, header_row, age_row


def export_to_word(institutes, filename="گزارش_دوره‌ها.docx"):
    """
    ایجاد فایل Word با اطلاعات مؤسسات و دوره‌ها.
    """
    doc = Document()
    
    # عنوان اصلی
    title = doc.add_heading('گزارش دوره‌های آموزشی تابستان ۱۴۰۵', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for inst in institutes:
        # نام مؤسسه
        doc.add_heading(f"🏛️ {inst['name']}", level=2)
        
        # مشخصات
        details = []
        if inst['discount']:
            details.append(f"تخفیف: {inst['discount']}")
        if inst['address']:
            details.append(f"آدرس: {inst['address']}")
        if inst['phone']:
            details.append(f"تلفن: {inst['phone']}")
        if details:
            doc.add_paragraph(" | ".join(details))

        # جدول دوره‌ها
        if inst['courses']:
            # ایجاد جدول با ۳ ستون
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'مقطع'
            hdr_cells[1].text = 'دوره'
            hdr_cells[2].text = 'قیمت'

            for course in inst['courses']:
                row_cells = table.add_row().cells
                row_cells[0].text = course['grade']
                row_cells[1].text = course['course']
                row_cells[2].text = course['price_text'] if course['price_text'] else 'توافقی'

        doc.add_paragraph()  # فاصله بین مؤسسات

    doc.save(filename)
    print(f"✅ فایل Word با نام '{filename}' ذخیره شد.")


def display_institutes(institutes):
    """نمایش در ترمینال (همان کد قبل)"""
    for inst in institutes:
        print(f"\n🏛️ مؤسسه: {inst['name']}")
        if inst['discount']:
            print(f"   تخفیف: {inst['discount']}")
        if inst['address']:
            print(f"   آدرس: {inst['address']}")
        if inst['phone']:
            print(f"   تلفن: {inst['phone']}")
        print("   دوره‌ها:")

        data = []
        for c in inst['courses']:
            data.append([c['grade'], c['course'], c['price_text'] or 'توافقی'])
        df = pd.DataFrame(data, columns=['مقطع', 'دوره', 'قیمت'])
        print(df.to_string(index=False))
        print("-" * 80)


if __name__ == "__main__":
    file_path = "برنامه دوره های آموزشی تابستان 1405-PART1.xlsx"  # مسیر فایل اکسل
    institutes, header, age = parse_excel(file_path)
    
    # نمایش در ترمینال
    display_institutes(institutes)
    
    # ایجاد فایل Word
    export_to_word(institutes)