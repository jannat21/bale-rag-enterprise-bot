import os
import re
from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_filename(name):
    main_name = re.split(r'[(\n]', name)[0].strip()
    forbidden_chars = r'[<>:"/\\|?*]'
    clean = re.sub(forbidden_chars, '_', main_name)
    clean = re.sub(r'\s+', ' ', clean).strip()
    if len(clean) > 50:
        clean = clean[:50].strip()
    return clean

def clean_price(text):
    if not text:
        return "قیمت موجود نیست"
    cleaned = text.replace('.', '').replace(',', '')
    cleaned = ' '.join(cleaned.split())
    return cleaned

def extract_courses_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    sheet = wb.active

    # اطلاعات سرستون‌ها (مقطع و سن)
    col_info = {}
    for col_idx in range(2, 8):
        level = sheet.cell(row=1, column=col_idx).value or ""
        age = sheet.cell(row=2, column=col_idx).value or ""
        col_info[col_idx] = {"level": str(level).strip(), "age": str(age).strip()}

    # خواندن تمام سطرها
    rows = []
    for r in sheet.iter_rows(min_row=1, values_only=True):
        rows.append(list(r))

    centers = []
    i = 2  # شروع از سطر سوم (داده‌ها)

    while i < len(rows):
        # پیدا کردن شروع موسسه بعدی
        if rows[i][0] is not None and str(rows[i][0]).strip() != "":
            center_name = str(rows[i][0]).strip()
            i += 1

            # جمع‌آوری تمام سلول‌های غیرخالی در ستون‌های B تا G برای این موسسه
            # به‌صورت لیستی از (سطر, ستون, مقدار)
            cells = []
            while i < len(rows):
                # بررسی پایان موسسه (شروع موسسه جدید یا بخش یارانه)
                if rows[i][0] is not None and str(rows[i][0]).strip() != "":
                    break
                if rows[i][1] is not None and ("مبلغ" in str(rows[i][1]) or "خانواده" in str(rows[i][1])):
                    break

                # بررسی ستون‌های B تا G
                for col_idx in range(2, 8):
                    val = rows[i][col_idx - 1]
                    if val is not None and str(val).strip() != "":
                        cells.append({
                            "row": i,
                            "col": col_idx,
                            "value": str(val).strip()
                        })
                i += 1

            # پردازش سلول‌های جمع‌آوری شده برای هر ستون به‌صورت جداگانه
            center_courses = []
            for col_idx in range(2, 8):
                # فیلتر سلول‌های این ستون و مرتب‌سازی بر اساس سطر
                col_cells = [c for c in cells if c["col"] == col_idx]
                col_cells.sort(key=lambda x: x["row"])

                # استخراج مقادیر
                values = [c["value"] for c in col_cells]

                # جفت‌سازی عنوان و قیمت
                j = 0
                while j < len(values):
                    current = values[j]
                    next_val = values[j+1] if j+1 < len(values) else None

                    # اگر فعلی قیمت باشد و بعدی عنوان باشد → جفت (قیمت, عنوان)
                    if "تومان" in current and next_val is not None and "تومان" not in next_val:
                        title = next_val
                        price = clean_price(current)
                        center_courses.append({
                            "title": title,
                            "price": price,
                            "level": col_info[col_idx]["level"],
                            "age": col_info[col_idx]["age"]
                        })
                        j += 2
                    # اگر فعلی عنوان باشد و بعدی قیمت باشد → جفت (عنوان, قیمت)
                    elif "تومان" not in current and next_val is not None and "تومان" in next_val:
                        title = current
                        price = clean_price(next_val)
                        center_courses.append({
                            "title": title,
                            "price": price,
                            "level": col_info[col_idx]["level"],
                            "age": col_info[col_idx]["age"]
                        })
                        j += 2
                    # اگر فعلی عنوان باشد و بعدی قیمت نباشد (یا پایان لیست)
                    elif "تومان" not in current:
                        # عنوان بدون قیمت
                        center_courses.append({
                            "title": current,
                            "price": "قیمت موجود نیست",
                            "level": col_info[col_idx]["level"],
                            "age": col_info[col_idx]["age"]
                        })
                        j += 1
                    # اگر فعلی قیمت باشد و بعدی قیمت یا پایان لیست
                    else:
                        # قیمت تنها - نادیده گرفته می‌شود
                        j += 1

            centers.append({
                "name": center_name,
                "courses": center_courses
            })
        else:
            i += 1

    return centers

def create_word_files(centers, output_dir="output_docs"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for center in centers:
        filename = clean_filename(center["name"])
        filepath = os.path.join(output_dir, f"{filename}.docx")

        doc = Document()
        heading = doc.add_heading(center["name"], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("لیست دوره‌های آموزشی", level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "عنوان دوره"
        hdr[1].text = "مبلغ (تومان)"
        hdr[2].text = "مقطع تحصیلی"
        hdr[3].text = "بازه سنی"

        for course in center["courses"]:
            row = table.add_row().cells
            row[0].text = course["title"]
            row[1].text = course["price"]
            row[2].text = course["level"]
            row[3].text = course["age"]

        doc.save(filepath)
        print(f"✅ {filename}.docx ایجاد شد (تعداد دوره‌ها: {len(center['courses'])}")

if __name__ == "__main__":
    excel_file = "part2.xlsx"
    centers = extract_courses_from_excel(excel_file)
    create_word_files(centers)
    print("🎯 همه فایل‌ها با موفقیت تولید شدند.")