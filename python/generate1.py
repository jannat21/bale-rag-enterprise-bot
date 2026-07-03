import os
import re
from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clean_filename(name):
    main_name = re.split(r'[(\n]', name)[0].strip()
    forbidden_chars = r'[<>:"/\\|?*]'
    clean = re.sub(forbidden_chars, '_', main_name)
    clean = re.sub(r'\s+', ' ', clean).strip()
    if len(clean) > 50:
        clean = clean[:50].strip()
    return clean

def clean_price(text):
    """حذف نقطه و کاما از رشته قیمت و یکسان‌سازی فاصله‌ها"""
    if not text:
        return "قیمت موجود نیست"
    # حذف نقطه و کاما
    cleaned = text.replace('.', '').replace(',', '')
    # حذف فاصله‌های اضافی
    cleaned = ' '.join(cleaned.split())
    return cleaned

def extract_courses_from_excel(file_path):
    wb = load_workbook(file_path, data_only=True)
    sheet = wb.active

    col_info = {}
    for col_idx in range(2, 8):
        level = sheet.cell(row=1, column=col_idx).value or ""
        age = sheet.cell(row=2, column=col_idx).value or ""
        col_info[col_idx] = {"level": str(level).strip(), "age": str(age).strip()}

    centers = []
    current_center = None
    row_idx = 3

    while row_idx <= sheet.max_row:
        row = sheet[row_idx]
        center_name_cell = row[0]

        if center_name_cell.value is not None and str(center_name_cell.value).strip() != "":
            if current_center is not None:
                centers.append(current_center)
            current_center = {
                "name": str(center_name_cell.value).strip(),
                "courses": []
            }
            row_idx += 1
            continue

        if current_center is None:
            break

        # بررسی پایان بخش یارانه
        if row[0].value is None and row[1].value is not None:
            if "مبلغ" in str(row[1].value) or "خانواده" in str(row[1].value):
                break

        titles = {}
        has_title = False
        for col_idx in range(2, 8):
            val = row[col_idx - 1].value
            if val is not None and str(val).strip() != "":
                text = str(val).strip()
                if "تومان" not in text:
                    titles[col_idx] = text
                    has_title = True

        if not has_title:
            row_idx += 1
            continue

        prices = {}
        next_row = sheet[row_idx + 1] if row_idx + 1 <= sheet.max_row else None
        if next_row is not None:
            is_price_row = False
            for col_idx in range(2, 8):
                val = next_row[col_idx - 1].value
                if val is not None and "تومان" in str(val):
                    is_price_row = True
                    break
            if is_price_row:
                for col_idx in range(2, 8):
                    val = next_row[col_idx - 1].value
                    if val is not None and str(val).strip() != "":
                        prices[col_idx] = clean_price(str(val).strip())   # ← اصلاح قیمت
                row_idx += 1

        for col_idx, title in titles.items():
            price = prices.get(col_idx, "قیمت موجود نیست")
            level = col_info[col_idx]["level"]
            age = col_info[col_idx]["age"]
            current_center["courses"].append({
                "title": title,
                "price": price,
                "level": level,
                "age": age
            })

        row_idx += 1

    if current_center is not None:
        centers.append(current_center)

    return centers


def create_word_files(centers, output_dir="output_docs"):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    for center in centers:
        filename = clean_filename(center["name"])
        filepath = os.path.join(output_dir, f"{filename}.docx")

        doc = Document()
        heading = doc.add_heading(center["name"], level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("لیست دوره‌های آموزشی", level=2)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "عنوان دوره"
        hdr[1].text = "مبلغ (تومان)"
        hdr[2].text = "مقطع تحصیلی"
        hdr[3].text = "بازه سنی"

        for course in center["courses"]:
            row = table.add_row().cells
            row[0].text = course["title"]
            row[1].text = course["price"]
            row[2].text = course["level"]
            row[3].text = course["age"]

        doc.save(filepath)
        print(f"✅ {filename}.docx ایجاد شد (تعداد دوره‌ها: {len(center['courses'])}")


if __name__ == "__main__":
    excel_file = "part2.xlsx"   # مسیر فایل Excel
    centers = extract_courses_from_excel(excel_file)
    create_word_files(centers)
    print("🎯 همه فایل‌ها با موفقیت تولید شدند.")